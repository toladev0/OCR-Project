# app.py
import streamlit as st
from utils.preprocessing import preprocess_image
from utils.ocr import run_ocr
from utils.pdf_utils import pdf_to_images
from PIL import Image
import numpy as np
import concurrent.futures
from config import TESSERACT_PATH, DEFAULT_LANGS
import pytesseract
import requests
from io import BytesIO

# Export dependencies
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT
import tempfile
import os

pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

# ── Streamlit UI ────────────────────────────────────────────────────────────
st.set_page_config(page_title="Khmer OCR", page_icon="assets/khmerOCR_logo.png", layout="wide")

st.image("assets/khmerOCR_logo.png", width=100)
st.markdown('<div style="font-size:2em; font-weight:bold;">Khmer OCR</div>', unsafe_allow_html=True)
st.markdown("Extract text from images or PDFs instantly")

# ── Sidebar ──────────────────────────────────────────────────────────────────
with st.sidebar:
    st.title("⚙️ Settings")
    langs = ["eng", "khm"]
    selected_langs = st.multiselect("Select OCR Languages", langs, default=DEFAULT_LANGS)
    lang_str = "+".join(selected_langs)

    preprocess_option = st.selectbox(
        "Choose Preprocessing Method",
        ["Auto", "None", "Grayscale", "Threshold", "Adaptive Threshold"],
        index=1
    )

    show_boxes = st.checkbox("Show Detected Text Boxes", value=False)
    st.info("💡 Drag & drop multiple files below or click to upload.")


# ── Helper: fetch image from URL ─────────────────────────────────────────────
def load_image_from_url(url: str):
    """Download an image from a URL and return a PIL Image, or None on failure."""
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        response = requests.get(url.strip(), headers=headers, timeout=10)
        response.raise_for_status()
        content_type = response.headers.get("Content-Type", "")
        if "image" not in content_type:
            st.error(f"URL does not point to an image (Content-Type: {content_type})")
            return None
        return Image.open(BytesIO(response.content)).convert("RGB")
    except requests.exceptions.MissingSchema:
        st.error("Invalid URL — please include http:// or https://")
    except requests.exceptions.ConnectionError:
        st.error("Could not connect. Check the URL and your internet connection.")
    except requests.exceptions.Timeout:
        st.error("Request timed out. The server took too long to respond.")
    except Exception as e:
        st.error(f"Failed to load image from URL: {e}")
    return None


# ── Helper: export to Word (.docx) ───────────────────────────────────────────
def build_docx(source_name: str, pages_text: list[str]) -> bytes:
    """Create a .docx file from a list of page texts and return as bytes."""
    doc = DocxDocument()

    # Title
    title = doc.add_heading(f"OCR Results — {source_name}", level=1)
    title.runs[0].font.size = Pt(16)

    for idx, text in enumerate(pages_text, start=1):
        doc.add_heading(f"Page {idx}", level=2)
        # Preserve line breaks
        for line in text.splitlines():
            para = doc.add_paragraph(line)
            para.runs[0].font.size = Pt(11) if para.runs else None
        if idx < len(pages_text):
            doc.add_page_break()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Helper: export to PDF ─────────────────────────────────────────────────────
def build_pdf(source_name: str, pages_text: list[str]) -> bytes:
    """Create a PDF file from a list of page texts and return as bytes."""
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            leftMargin=72, rightMargin=72,
                            topMargin=72, bottomMargin=72)

    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    heading_style = styles["Heading2"]
    body_style = ParagraphStyle(
        "KhmerBody",
        parent=styles["Normal"],
        fontSize=11,
        leading=16,
        alignment=TA_LEFT,
        fontName="Helvetica",
    )

    story = []
    story.append(Paragraph(f"OCR Results — {source_name}", title_style))
    story.append(Spacer(1, 12))

    for idx, text in enumerate(pages_text, start=1):
        story.append(Paragraph(f"Page {idx}", heading_style))
        story.append(Spacer(1, 6))
        # Escape XML special characters and preserve line breaks
        safe_text = (text
                     .replace("&", "&amp;")
                     .replace("<", "&lt;")
                     .replace(">", "&gt;"))
        for line in safe_text.splitlines():
            story.append(Paragraph(line or " ", body_style))
        if idx < len(pages_text):
            story.append(PageBreak())

    doc.build(story)
    return buf.getvalue()


# ── Helper: process a single page ───────────────────────────────────────────
def process_page(file_idx, page_idx, img, lang_str, preprocess_option, show_boxes, local_cache):
    key = f"{file_idx}_{page_idx}"
    if key in local_cache:
        processed_img = local_cache[key]
    else:
        processed_img = preprocess_image(img, preprocess_option)
        local_cache[key] = processed_img
    text, display_img = run_ocr(processed_img, lang_str, draw_boxes=show_boxes)
    return text, display_img


# ── Helper: render results for a source (file or URL) ───────────────────────
def render_source(source_label: str, file_idx: int, images: list):
    """Run OCR on all images, display results, and offer exports."""
    combined_text = [None] * len(images)
    display_images = [None] * len(images)
    local_cache = {}

    with concurrent.futures.ThreadPoolExecutor() as executor:
        futures = {
            executor.submit(
                process_page, file_idx, idx, np.array(img),
                lang_str, preprocess_option, show_boxes, local_cache
            ): idx
            for idx, img in enumerate(images)
        }
        for future in concurrent.futures.as_completed(futures):
            idx = futures[future]
            try:
                text, display_img = future.result()
                combined_text[idx] = text
                display_images[idx] = display_img
            except Exception as e:
                st.error(f"Error processing page {idx + 1}: {e}")

    # Per-page display
    for idx, img in enumerate(images):
        if combined_text[idx] is None:
            continue
        with st.expander(f"Page {idx + 1}", expanded=False):
            col1, col2 = st.columns([1, 1])
            with col1:
                show_original = st.checkbox("Show Original", key=f"orig_{file_idx}_{idx}")
                st.image(np.array(img) if show_original else display_images[idx], width=600)
            with col2:
                st.text_area("OCR Text", combined_text[idx], height=300,
                             key=f"text_{file_idx}_{idx}")
                st.download_button(
                    f"⬇️ Download Page {idx + 1} Text",
                    combined_text[idx],
                    file_name=f"ocr_page_{file_idx + 1}_{idx + 1}.txt",
                    key=f"dl_txt_{file_idx}_{idx}",
                )

    # Combined text for this source
    pages_with_text = [t for t in combined_text if t]
    if not pages_with_text:
        return

    file_text = "\n\n".join(pages_with_text)

    st.markdown("#### 📤 Export Results")
    col_txt, col_docx, col_pdf = st.columns(3)

    with col_txt:
        st.download_button(
            "⬇️ Download as .txt",
            file_text,
            file_name=f"ocr_{source_label}.txt",
            key=f"export_txt_{file_idx}",
        )

    with col_docx:
        docx_bytes = build_docx(source_label, [t for t in combined_text if t])
        st.download_button(
            "📄 Download as Word (.docx)",
            docx_bytes,
            file_name=f"ocr_{source_label}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"export_docx_{file_idx}",
        )

    with col_pdf:
        pdf_bytes = build_pdf(source_label, [t for t in combined_text if t])
        st.download_button(
            "📕 Download as PDF",
            pdf_bytes,
            file_name=f"ocr_{source_label}.pdf",
            mime="application/pdf",
            key=f"export_pdf_{file_idx}",
        )

    return file_text


# ── Input section ─────────────────────────────────────────────────────────────
tab_file, tab_url = st.tabs(["📁 Upload Files", "🔗 Image URL"])

all_files_text = []
file_idx_counter = 0  # global counter so keys never clash across tabs

# ── Tab 1: File Upload ────────────────────────────────────────────────────────
with tab_file:
    uploaded_files = st.file_uploader(
        "Upload Images or PDFs", type=["png", "jpg", "jpeg", "pdf"],
        accept_multiple_files=True,
    )

    if uploaded_files:
        for uploaded_file in uploaded_files:
            st.markdown(f"### 📁 {uploaded_file.name}")
            try:
                if uploaded_file.type == "application/pdf":
                    images = pdf_to_images(uploaded_file.read())
                else:
                    images = [Image.open(uploaded_file)]

                if images:
                    safe_name = uploaded_file.name.replace(" ", "_")
                    result = render_source(safe_name, file_idx_counter, images)
                    if result:
                        all_files_text.append(result)
            except Exception as e:
                st.error(f"Unexpected error with {uploaded_file.name}: {e}")
            file_idx_counter += 1

# ── Tab 2: URL Input ───────────────────────────────────────────────────────────
with tab_url:
    st.markdown("Paste one or more image URLs (one per line):")
    url_input = st.text_area("Image URLs", placeholder="https://example.com/image.jpg", height=120)

    if st.button("🔍 Run OCR on URLs"):
        urls = [u.strip() for u in url_input.splitlines() if u.strip()]
        if not urls:
            st.warning("Please enter at least one URL.")
        else:
            for url in urls:
                st.markdown(f"### 🔗 {url}")
                img = load_image_from_url(url)
                if img is not None:
                    safe_name = url.split("/")[-1].split("?")[0] or "url_image"
                    result = render_source(safe_name, file_idx_counter, [img])
                    if result:
                        all_files_text.append(result)
                file_idx_counter += 1

# ── Combined download for all sources ────────────────────────────────────────
if all_files_text:
    st.markdown("---")
    st.markdown("### 📦 Combined Export (All Sources)")
    all_text = "\n\n".join(all_files_text)

    col1, col2, col3 = st.columns(3)
    with col1:
        st.download_button("⬇️ All Files — .txt", all_text,
                           file_name="ocr_all_combined.txt")
    with col2:
        docx_all = build_docx("All Sources", all_files_text)
        st.download_button("📄 All Files — .docx", docx_all,
                           file_name="ocr_all_combined.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with col3:
        pdf_all = build_pdf("All Sources", all_files_text)
        st.download_button("📕 All Files — .pdf", pdf_all,
                           file_name="ocr_all_combined.pdf",
                           mime="application/pdf")
