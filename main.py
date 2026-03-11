import streamlit as st
import pytesseract
from PIL import Image
import io
from docx import Document
import base64

# Tesseract path
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Streamlit UI
st.title("Khmer OCR Project Demo")
st.write("Image → Khmer Text → Editable DOCX")

# Upload
uploaded_file = st.file_uploader("Choose Khmer image", type=['png','jpg','jpeg'])

if uploaded_file is not None:
    image = Image.open(uploaded_file)
    st.image(image, caption="Uploaded", width=400)
    
    # Process button
    if st.button("🔍 Extract Khmer Text"):
        with st.spinner("Running Tesseract khm..."):
            # Simple OCR
            text = pytesseract.image_to_string(image, lang='khm', config='--oem 1 --psm 6')
            
            # Show result
            st.subheader("Extracted Text")
            st.write(text)
            
            # DOCX download
            doc = Document()
            doc.add_heading("Khmer OCR Result", 0)
            doc.add_paragraph(text)
            
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            # Download button
            b64 = base64.b64encode(doc_bytes.read()).decode()
            href = f'<a href="data:application/octet-stream;base64,{b64}" download="khmer_result.docx">⬇️ Download DOCX</a>'
            st.markdown(href, unsafe_allow_html=True)
